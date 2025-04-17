import fitz  # PyMuPDF
import docx
import tempfile

def extract_text(filename: str, content_bytes: bytes) -> dict:
    with tempfile.NamedTemporaryFile(delete=False, suffix=filename[-4:]) as tmp:
        tmp.write(content_bytes)
        tmp_path = tmp.name

    if filename.endswith(".pdf"):
        return _parse_pdf(tmp_path)
    elif filename.endswith(".docx"):
        return _parse_docx(tmp_path)

    raise ValueError("Unsupported file type.")

def _parse_pdf(path: str) -> dict:
    doc = fitz.open(path)
    full_text = []
    useful_lines = 0

    for page in doc:
        page_text = page.get_text()
        full_text.append(page_text)
        useful_lines += sum(1 for line in page_text.splitlines() if line.strip())

    full_text_str = "\n".join(full_text)
    useful_ratio = useful_lines / len(full_text_str.splitlines()) if full_text_str else 0
    return {
        "text": full_text_str,
        "num_pages": len(doc),
        "useful_text_ratio": round(useful_ratio, 2)  # ✅ تم التعديل هنا
    }

def _parse_docx(path: str) -> dict:
    doc = docx.Document(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text_str = "\n".join(lines)
    useful_ratio = len(lines) / len(doc.paragraphs) if doc.paragraphs else 0
    return {
        "text": full_text_str,
        "num_pages": 1,
        "useful_text_ratio": round(useful_ratio, 2)  # ✅ وتم التعديل هنا كمان
    }
